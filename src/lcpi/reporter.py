from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn
from rich.console import Console
from rich.text import Text
import time
import subprocess
import json
import pathlib
import sys
import os
from datetime import datetime
import csv
from typing import List, Dict, Any, Optional
import hashlib
import pickle
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

# Nouvelles dépendances pour les améliorations
try:
    from jinja2 import Environment, FileSystemLoader, Template
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False
    print("Warning: Jinja2 non installé. Les templates ne seront pas disponibles.")

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Backend non-interactif
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("Warning: Matplotlib non installé. Les graphiques ne seront pas générés.")

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx non installé. L'export DOCX ne sera pas disponible.")

console = Console()

class ReportCache:
    """Système de cache pour les résultats de calculs."""
    
    def __init__(self, cache_dir: str = "cache"):
        self.cache_dir = pathlib.Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True)
        self.cache_file = self.cache_dir / "report_cache.pkl"
        self.cache_data = self._load_cache()
    
    def _load_cache(self) -> Dict[str, Any]:
        """Charge le cache depuis le fichier."""
        if self.cache_file.exists():
            try:
                with open(self.cache_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                console.print(f"[yellow]Impossible de charger le cache: {e}[/yellow]")
        return {}
    
    def _save_cache(self):
        """Sauvegarde le cache dans le fichier."""
        try:
            with open(self.cache_file, 'wb') as f:
                pickle.dump(self.cache_data, f)
        except Exception as e:
            console.print(f"[yellow]Impossible de sauvegarder le cache: {e}[/yellow]")
    
    def _calculate_file_hash(self, file_path: pathlib.Path) -> str:
        """Calcule le hash d'un fichier."""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception:
            return ""
    
    def get_cached_result(self, file_path: pathlib.Path) -> Optional[Dict[str, Any]]:
        """Récupère un résultat en cache si disponible."""
        file_hash = self._calculate_file_hash(file_path)
        cache_key = f"{file_path}_{file_hash}"
        
        if cache_key in self.cache_data:
            cached_data = self.cache_data[cache_key]
            # Vérifier que le cache n'est pas trop ancien (7 jours)
            if time.time() - cached_data.get('timestamp', 0) < 7 * 24 * 3600:
                return cached_data.get('result')
        return None
    
    def cache_result(self, file_path: pathlib.Path, result: Dict[str, Any]):
        """Met en cache un résultat."""
        file_hash = self._calculate_file_hash(file_path)
        cache_key = f"{file_path}_{file_hash}"
        
        self.cache_data[cache_key] = {
            'result': result,
            'timestamp': time.time(),
            'file_path': str(file_path)
        }
        self._save_cache()
    
    def clear_cache(self):
        """Vide le cache."""
        self.cache_data = {}
        if self.cache_file.exists():
            self.cache_file.unlink()
        console.print("[green]Cache vidé avec succès[/green]")

class ReportAnalyzer:
    """Analyseur intelligent pour les rapports de synthèse."""
    
    @staticmethod
    def generate_synthesis(results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Génère une synthèse intelligente des résultats."""
        synthesis = {
            'total_elements': len(results),
            'plugins_used': list(set(r.get('plugin', 'Inconnu') for r in results)),
            'status_summary': {},
            'critical_ratios': [],
            'warnings': [],
            'errors': [],
            'success_rate': 0.0
        }
        
        # Analyse des statuts
        for result in results:
            status = result.get('statut', 'Inconnu')
            synthesis['status_summary'][status] = synthesis['status_summary'].get(status, 0) + 1
        
        # Calcul du taux de succès
        success_count = synthesis['status_summary'].get('OK', 0)
        synthesis['success_rate'] = (success_count / len(results)) * 100 if results else 0
        
        # Analyse des ratios critiques
        for result in results:
            resultats = result.get('resultats', {})
            for key, value in resultats.items():
                if any(crit_key in key.lower() for crit_key in ['ratio', 'coefficient', 'securite', 'verification']):
                    try:
                        if isinstance(value, str):
                            # Extraire la valeur numérique si possible
                            import re
                            num_match = re.search(r'(\d+\.?\d*)', value)
                            if num_match:
                                num_value = float(num_match.group(1))
                            else:
                                continue
                        else:
                            num_value = float(value)
                        
                        synthesis['critical_ratios'].append({
                            'element': result.get('element_id', 'Inconnu'),
                            'parameter': key,
                            'value': num_value,
                            'status': result.get('statut', 'Inconnu')
                        })
                    except (ValueError, TypeError):
                        continue
        
        # Tri des ratios critiques par valeur
        synthesis['critical_ratios'].sort(key=lambda x: x['value'])
        
        # Collecte des avertissements et erreurs
        for result in results:
            if result.get('statut') == 'Avertissement':
                synthesis['warnings'].append({
                    'element': result.get('element_id', 'Inconnu'),
                    'plugin': result.get('plugin', 'Inconnu'),
                    'details': result.get('resultats', {})
                })
            elif result.get('statut') == 'Erreur':
                synthesis['errors'].append({
                    'element': result.get('element_id', 'Inconnu'),
                    'plugin': result.get('plugin', 'Inconnu'),
                    'details': result.get('resultats', {})
                })
        
        return synthesis
    
    @staticmethod
    def compare_reports(current_results: List[Dict[str, Any]], previous_results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Compare deux rapports et identifie les différences."""
        comparison = {
            'added_elements': [],
            'removed_elements': [],
            'modified_elements': [],
            'summary_changes': {}
        }
        
        # Créer des dictionnaires pour faciliter la comparaison
        current_dict = {r.get('element_id'): r for r in current_results}
        previous_dict = {r.get('element_id'): r for r in previous_results}
        
        # Éléments ajoutés
        for element_id in current_dict:
            if element_id not in previous_dict:
                comparison['added_elements'].append(element_id)
        
        # Éléments supprimés
        for element_id in previous_dict:
            if element_id not in current_dict:
                comparison['removed_elements'].append(element_id)
        
        # Éléments modifiés
        for element_id in current_dict:
            if element_id in previous_dict:
                current = current_dict[element_id]
                previous = previous_dict[element_id]
                
                changes = []
                
                # Comparer les statuts
                if current.get('statut') != previous.get('statut'):
                    changes.append({
                        'type': 'status',
                        'field': 'statut',
                        'old': previous.get('statut'),
                        'new': current.get('statut')
                    })
                
                # Comparer les résultats numériques
                current_resultats = current.get('resultats', {})
                previous_resultats = previous.get('resultats', {})
                
                for key in set(current_resultats.keys()) | set(previous_resultats.keys()):
                    current_val = current_resultats.get(key)
                    previous_val = previous_resultats.get(key)
                    
                    if current_val != previous_val:
                        # Essayer de calculer la différence pour les valeurs numériques
                        try:
                            if isinstance(current_val, str) and isinstance(previous_val, str):
                                import re
                                curr_num = re.search(r'(\d+\.?\d*)', current_val)
                                prev_num = re.search(r'(\d+\.?\d*)', previous_val)
                                
                                if curr_num and prev_num:
                                    curr_float = float(curr_num.group(1))
                                    prev_float = float(prev_num.group(1))
                                    if prev_float != 0:
                                        percentage_change = ((curr_float - prev_float) / prev_float) * 100
                                        changes.append({
                                            'type': 'numeric',
                                            'field': key,
                                            'old': previous_val,
                                            'new': current_val,
                                            'change_percent': percentage_change
                                        })
                                    else:
                                        changes.append({
                                            'type': 'value',
                                            'field': key,
                                            'old': previous_val,
                                            'new': current_val
                                        })
                                else:
                                    changes.append({
                                        'type': 'value',
                                        'field': key,
                                        'old': previous_val,
                                        'new': current_val
                                    })
    else:
                                changes.append({
                                    'type': 'value',
                                    'field': key,
                                    'old': previous_val,
                                    'new': current_val
                                })
                        except (ValueError, TypeError):
                            changes.append({
                                'type': 'value',
                                'field': key,
                                'old': previous_val,
                                'new': current_val
                            })
                
                if changes:
                    comparison['modified_elements'].append({
                        'element_id': element_id,
                        'changes': changes
                    })
        
        # Résumé des changements
        comparison['summary_changes'] = {
            'total_elements_current': len(current_results),
            'total_elements_previous': len(previous_results),
            'elements_added': len(comparison['added_elements']),
            'elements_removed': len(comparison['removed_elements']),
            'elements_modified': len(comparison['modified_elements'])
        }
        
        return comparison

class ReportGenerator:
    """Générateur de rapports amélioré avec support multi-formats et templates."""
    
    def __init__(self, project_dir: str, enable_cache: bool = True, max_workers: int = 4):
        self.project_dir = pathlib.Path(project_dir).resolve()
        self.template_dir = self.project_dir / "templates"
        self.output_dir = self.project_dir / "output"
        self.output_dir.mkdir(exist_ok=True)
        self.max_workers = max_workers
        
        # Système de cache
        self.cache = ReportCache() if enable_cache else None
        
        # Configuration Jinja2
        if JINJA2_AVAILABLE:
            self.jinja_env = Environment(
                loader=FileSystemLoader([str(self.template_dir), str(pathlib.Path(__file__).parent / "templates")]),
                autoescape=True
            )
    
    def _process_single_file(self, yml_file: pathlib.Path, plugin_name: str, command: str) -> Optional[Dict[str, Any]]:
        """Traite un seul fichier YML et retourne le résultat."""
        try:
            # Vérifier le cache si activé
            if self.cache:
                cached_result = self.cache.get_cached_result(yml_file)
                if cached_result:
                    return cached_result
            
            script_path = pathlib.Path(__file__).parent.parent / "main.py"
            cmd = [sys.executable, str(script_path), plugin_name, command, str(yml_file), "--json"]
            
                process = subprocess.run(
                    cmd, capture_output=True, text=True, check=False, encoding='utf-8', errors='ignore'
                )
                
                if process.returncode != 0:
                    raise subprocess.CalledProcessError(
                        returncode=process.returncode,
                        cmd=cmd,
                        stderr=process.stderr
                    )

                output = process.stdout
                start = output.find('{')
                end = output.rfind('}') + 1
                if start != -1 and end != 0:
                    json_output = output[start:end]
                    data = json.loads(json_output)
                
                # Mettre en cache le résultat
                if self.cache:
                    self.cache.cache_result(yml_file, data)
                
                return data
                else:
                        console.log(f"[yellow]Avertissement[/yellow]: Pas de sortie JSON pour {yml_file.name}.")
                return None

            except subprocess.CalledProcessError as e:
                    console.print(f"[bold red]Erreur lors de l'analyse de {yml_file.name}[/bold red]")
                    console.print(f"[red]  Code de sortie: {e.returncode}[/red]")
                    error_output = e.stderr or e.stdout or ""
                    console.print(f"[red]  Erreur: {error_output.strip()}[/red]")
            return None
            except json.JSONDecodeError:
                    console.print(f"[bold red]Erreur de décodage JSON pour {yml_file.name}[/bold red]")
            return None
            except Exception as e:
                    console.print(f"[bold red]Une erreur inattendue est survenue avec {yml_file.name}: {e}[/bold red]")
            return None
    
    def analyze_project_parallel(self, plugin_commands: Dict[str, str]) -> List[Dict[str, Any]]:
        """Analyse le projet en parallèle avec cache."""
        results = []
        files_to_process = []
        
        # Collecter tous les fichiers à traiter
        for plugin in plugin_commands.keys():
            plugin_dir = pathlib.Path(__file__).parent.parent / plugin
            if plugin_dir.is_dir():
                for yml_file in plugin_dir.glob("**/*.yml"):
                    try:
                        plugin_name = yml_file.relative_to(pathlib.Path(__file__).parent.parent).parts[0]
                        command = plugin_commands.get(plugin_name)
                        if command:
                            files_to_process.append((yml_file, plugin_name, command))
                    except ValueError:
                        continue
        
        if not files_to_process:
            console.print("[yellow]Aucun fichier .yml à analyser n'a été trouvé dans les plugins.[/yellow]")
            return results
        
        # Traitement parallèle
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Soumettre toutes les tâches
            future_to_file = {
                executor.submit(self._process_single_file, yml_file, plugin_name, command): yml_file
                for yml_file, plugin_name, command in files_to_process
            }
            
            # Barre de progression
            with Progress(
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                SpinnerColumn(),
                console=console,
                transient=True
            ) as progress:
                task = progress.add_task("[green]Analyse parallèle...", total=len(files_to_process))
                
                # Traiter les résultats au fur et à mesure
                for future in as_completed(future_to_file):
                    yml_file = future_to_file[future]
                    try:
                        result = future.result()
                        if result:
                            results.append(result)
                    except Exception as e:
                        console.print(f"[red]Erreur avec {yml_file.name}: {e}[/red]")
                    
                progress.advance(task)
        
        return results
    
    def generate_graphs(self, results: List[Dict[str, Any]]) -> List[str]:
        """Génère des graphiques à partir des résultats et retourne les chemins des images."""
        if not MATPLOTLIB_AVAILABLE:
            return []
        
        graph_paths = []
        
        # Graphique 1: Répartition par plugin
        plugin_counts = {}
        for result in results:
            plugin = result.get('plugin', 'Inconnu')
            plugin_counts[plugin] = plugin_counts.get(plugin, 0) + 1
        
        if plugin_counts:
            plt.figure(figsize=(10, 6))
            plt.pie(plugin_counts.values(), labels=plugin_counts.keys(), autopct='%1.1f%%')
            plt.title('Répartition des analyses par plugin')
            graph_path = str(self.output_dir / "repartition_plugins.png")
            plt.savefig(graph_path, dpi=300, bbox_inches='tight')
            plt.close()
            graph_paths.append(graph_path)
        
        # Graphique 2: Statuts des résultats
        status_counts = {}
        for result in results:
            status = result.get('statut', 'Inconnu')
            status_counts[status] = status_counts.get(status, 0) + 1
        
        if status_counts:
            plt.figure(figsize=(10, 6))
            bars = plt.bar(status_counts.keys(), status_counts.values(), color=['green', 'red', 'orange'])
            plt.title('Statuts des résultats d\'analyse')
            plt.ylabel('Nombre d\'éléments')
            
            # Ajouter les valeurs sur les barres
            for bar in bars:
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                        f'{int(height)}', ha='center', va='bottom')
            
            graph_path = str(self.output_dir / "statuts_resultats.png")
            plt.savefig(graph_path, dpi=300, bbox_inches='tight')
            plt.close()
            graph_paths.append(graph_path)
        
        # Graphique 3: Ratios critiques (si disponibles)
        synthesis = ReportAnalyzer.generate_synthesis(results)
        if synthesis['critical_ratios']:
            plt.figure(figsize=(12, 8))
            
            # Prendre les 10 premiers ratios
            top_ratios = synthesis['critical_ratios'][:10]
            elements = [r['element'] for r in top_ratios]
            values = [r['value'] for r in top_ratios]
            colors_list = ['red' if r['status'] != 'OK' else 'green' for r in top_ratios]
            
            bars = plt.barh(elements, values, color=colors_list)
            plt.title('Top 10 des Ratios Critiques')
            plt.xlabel('Valeur du ratio')
            
            # Ajouter les valeurs sur les barres
            for i, bar in enumerate(bars):
                width = bar.get_width()
                plt.text(width + 0.01, bar.get_y() + bar.get_height()/2,
                        f'{values[i]:.2f}', ha='left', va='center')
            
            plt.tight_layout()
            graph_path = str(self.output_dir / "ratios_critiques.png")
            plt.savefig(graph_path, dpi=300, bbox_inches='tight')
            plt.close()
            graph_paths.append(graph_path)
        
        return graph_paths
    
    def generate_html_report(self, results: List[Dict[str, Any]], template_name: str = "default.html", 
                           synthesis: Optional[Dict[str, Any]] = None, 
                           comparison: Optional[Dict[str, Any]] = None) -> str:
        """Génère un rapport HTML avec template Jinja2."""
        if not JINJA2_AVAILABLE:
            console.print("[red]Jinja2 n'est pas installé. Impossible de générer le rapport HTML.[/red]")
            return ""
        
        # Générer la synthèse si non fournie
        if synthesis is None:
            synthesis = ReportAnalyzer.generate_synthesis(results)
        
        # Préparer les données pour le template
        template_data = {
            'project_name': self.project_dir.name,
            'generation_date': datetime.now().strftime("%d/%m/%Y à %H:%M"),
            'results': results,
            'total_elements': len(results),
            'plugins': list(set(r.get('plugin', 'Inconnu') for r in results)),
            'graphs': self.generate_graphs(results),
            'synthesis': synthesis,
            'comparison': comparison
        }
        
        try:
            template = self.jinja_env.get_template(template_name)
            html_content = template.render(**template_data)
            
            output_path = str(self.output_dir / "rapport_lcpi.html")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return output_path
        except Exception as e:
            console.print(f"[red]Erreur lors de la génération du rapport HTML: {e}[/red]")
            return ""
    
    def generate_docx_report(self, results: List[Dict[str, Any]], 
                           synthesis: Optional[Dict[str, Any]] = None,
                           comparison: Optional[Dict[str, Any]] = None) -> str:
        """Génère un rapport Word (DOCX)."""
        if not DOCX_AVAILABLE:
            console.print("[red]python-docx n'est pas installé. Impossible de générer le rapport DOCX.[/red]")
            return ""
        
        # Générer la synthèse si non fournie
        if synthesis is None:
            synthesis = ReportAnalyzer.generate_synthesis(results)
        
        doc = Document()
        
        # Titre
        doc.add_heading('Rapport d\'Analyse LCPI-CLI', 0)
        doc.add_paragraph(f'Généré le {datetime.now().strftime("%d/%m/%Y à %H:%M")}')
        doc.add_paragraph(f'Projet: {self.project_dir.name}')
        
        # Synthèse intelligente
        doc.add_heading('Synthèse Intelligente', level=1)
        doc.add_paragraph(f'Nombre total d\'éléments analysés: {synthesis["total_elements"]}')
        doc.add_paragraph(f'Taux de succès: {synthesis["success_rate"]:.1f}%')
        doc.add_paragraph(f'Plugins utilisés: {", ".join(synthesis["plugins_used"])}')
        
        # Tableau des statuts
        if synthesis['status_summary']:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Statut'
            hdr_cells[1].text = 'Nombre'
            
            for status, count in synthesis['status_summary'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = status
                row_cells[1].text = str(count)
        
        # Ratios critiques
        if synthesis['critical_ratios']:
            doc.add_heading('Ratios Critiques', level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Élément'
            hdr_cells[1].text = 'Paramètre'
            hdr_cells[2].text = 'Valeur'
            hdr_cells[3].text = 'Statut'
            
            for ratio in synthesis['critical_ratios'][:10]:  # Top 10
                row_cells = table.add_row().cells
                row_cells[0].text = ratio['element']
                row_cells[1].text = ratio['parameter']
                row_cells[2].text = str(ratio['value'])
                row_cells[3].text = ratio['status']
        
        # Comparaison si disponible
        if comparison:
            doc.add_heading('Comparaison avec Rapport Précédent', level=1)
            doc.add_paragraph(f'Éléments ajoutés: {comparison["summary_changes"]["elements_added"]}')
            doc.add_paragraph(f'Éléments supprimés: {comparison["summary_changes"]["elements_removed"]}')
            doc.add_paragraph(f'Éléments modifiés: {comparison["summary_changes"]["elements_modified"]}')
            
            if comparison['modified_elements']:
                doc.add_heading('Modifications Détailées', level=2)
                for mod in comparison['modified_elements'][:5]:  # Top 5
                    doc.add_paragraph(f'Élément: {mod["element_id"]}', style='Heading3')
                    for change in mod['changes']:
                        if change['type'] == 'numeric' and 'change_percent' in change:
                            doc.add_paragraph(f'  {change["field"]}: {change["old"]} → {change["new"]} ({change["change_percent"]:+.1f}%)')
                        else:
                            doc.add_paragraph(f'  {change["field"]}: {change["old"]} → {change["new"]}')
        
        # Graphiques
        graph_paths = self.generate_graphs(results)
        if graph_paths:
            doc.add_heading('Graphiques', level=1)
            for graph_path in graph_paths:
                doc.add_picture(graph_path, width=Inches(6))
                doc.add_paragraph()
        
        # Détails par élément
        doc.add_heading('Détails des analyses', level=1)
        
        for result in results:
            element_id = result.get('element_id', 'Inconnu')
            plugin = result.get('plugin', 'Inconnu')
            statut = result.get('statut', 'Inconnu')
            
            doc.add_heading(f'Élément: {element_id}', level=2)
            doc.add_paragraph(f'Plugin: {plugin}')
            doc.add_paragraph(f'Statut: {statut}')
            
            # Résultats détaillés
            resultats = result.get('resultats', {})
            if resultats:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Paramètre'
                hdr_cells[1].text = 'Valeur'
                
                for key, value in resultats.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = key.replace('_', ' ').title()
                    row_cells[1].text = str(value)
            
            doc.add_paragraph()
        
        output_path = str(self.output_dir / "rapport_lcpi.docx")
        doc.save(output_path)
        return output_path
    
    def generate_csv_report(self, results: List[Dict[str, Any]]) -> str:
        """Génère un rapport CSV pour analyse dans un tableur."""
        output_path = str(self.output_dir / "rapport_lcpi.csv")
        
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['element_id', 'plugin', 'statut', 'parametre', 'valeur']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            
            writer.writeheader()
            for result in results:
                element_id = result.get('element_id', 'Inconnu')
                plugin = result.get('plugin', 'Inconnu')
                statut = result.get('statut', 'Inconnu')
                
                resultats = result.get('resultats', {})
                if resultats:
                    for key, value in resultats.items():
                        writer.writerow({
                            'element_id': element_id,
                            'plugin': plugin,
                            'statut': statut,
                            'parametre': key.replace('_', ' ').title(),
                            'valeur': str(value)
                        })
                else:
                    writer.writerow({
                        'element_id': element_id,
                        'plugin': plugin,
                        'statut': statut,
                        'parametre': '',
                        'valeur': ''
                    })
        
        return output_path
    
    def generate_pdf_report(self, results: List[Dict[str, Any]], template_name: str = "default.html",
                          synthesis: Optional[Dict[str, Any]] = None,
                          comparison: Optional[Dict[str, Any]] = None) -> str:
        """Génère un rapport PDF amélioré avec graphiques et synthèse."""
        # Générer la synthèse si non fournie
        if synthesis is None:
            synthesis = ReportAnalyzer.generate_synthesis(results)
        
        doc = SimpleDocTemplate(str(self.output_dir / "rapport_lcpi.pdf"), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Titre
        story.append(Paragraph("Rapport d'Analyse LCPI-CLI", styles['Title']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Projet: {self.project_dir.name}", styles['Normal']))
        story.append(Paragraph(f"Généré le: {datetime.now().strftime('%d/%m/%Y à %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Synthèse intelligente
        story.append(Paragraph("Synthèse Intelligente", styles['Heading1']))
        story.append(Paragraph(f"Nombre total d'éléments analysés: {synthesis['total_elements']}", styles['Normal']))
        story.append(Paragraph(f"Taux de succès: {synthesis['success_rate']:.1f}%", styles['Normal']))
        story.append(Paragraph(f"Plugins utilisés: {', '.join(synthesis['plugins_used'])}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Tableau des statuts
        if synthesis['status_summary']:
            story.append(Paragraph("Répartition par statut:", styles['Heading2']))
            status_data = [['Statut', 'Nombre']]
            for status, count in synthesis['status_summary'].items():
                status_data.append([status, str(count)])
            
            status_table = Table(status_data)
            status_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(status_table)
            story.append(Spacer(1, 12))
        
        # Ratios critiques
        if synthesis['critical_ratios']:
            story.append(Paragraph("Ratios Critiques (Top 10):", styles['Heading2']))
            ratio_data = [['Élément', 'Paramètre', 'Valeur', 'Statut']]
            for ratio in synthesis['critical_ratios'][:10]:
                ratio_data.append([
                    ratio['element'],
                    ratio['parameter'],
                    f"{ratio['value']:.2f}",
                    ratio['status']
                ])
            
            ratio_table = Table(ratio_data)
            ratio_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(ratio_table)
            story.append(Spacer(1, 12))
        
        # Comparaison si disponible
        if comparison:
            story.append(Paragraph("Comparaison avec Rapport Précédent", styles['Heading1']))
            story.append(Paragraph(f"Éléments ajoutés: {comparison['summary_changes']['elements_added']}", styles['Normal']))
            story.append(Paragraph(f"Éléments supprimés: {comparison['summary_changes']['elements_removed']}", styles['Normal']))
            story.append(Paragraph(f"Éléments modifiés: {comparison['summary_changes']['elements_modified']}", styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Graphiques
        graph_paths = self.generate_graphs(results)
        if graph_paths:
            story.append(Paragraph("Graphiques", styles['Heading1']))
            for graph_path in graph_paths:
                try:
                    img = Image(graph_path, width=6*inch, height=4*inch)
                    story.append(img)
                    story.append(Spacer(1, 12))
                except Exception as e:
                    console.print(f"[yellow]Impossible d'inclure le graphique {graph_path}: {e}[/yellow]")
        
        # Détails par élément
        story.append(Paragraph("Détails des analyses", styles['Heading1']))
        
        for result in results:
            element_id = result.get('element_id', 'Inconnu')
            plugin = result.get('plugin', 'Inconnu')
            statut = result.get('statut', 'Inconnu')
            
            story.append(Paragraph(f"Élément: {element_id}", styles['Heading2']))
            story.append(Paragraph(f"Plugin: {plugin} | Statut: {statut}", styles['Normal']))
            
            result_data = result.get('resultats', {})
            for key, value in result_data.items():
                text = f"<b>{key.replace('_', ' ').title()}:</b> {value}"
                story.append(Paragraph(text, styles['Normal']))
            
            story.append(Spacer(1, 12))
        
        with console.status("[bold green]Génération du rapport PDF...", spinner="dots") as status:
            doc.build(story)
            time.sleep(1)
        
        return str(self.output_dir / "rapport_lcpi.pdf")

def generate_pdf_report(results_list: list, output_filename: str):
    """Fonction legacy pour compatibilité."""
    generator = ReportGenerator(pathlib.Path(output_filename).parent)
    return generator.generate_pdf_report(results_list)

def run_analysis_and_generate_report(project_dir: str, output_format: str = "pdf", template_name: str = "default.html",
                                   compare_with: Optional[str] = None, enable_cache: bool = True, max_workers: int = 4):
    """
    Fonction principale améliorée pour scanner le projet, exécuter les calculs
    et générer le rapport final avec support multi-formats et parallélisation.
    """
    generator = ReportGenerator(project_dir, enable_cache, max_workers)
    results = []
    
    plugin_commands = {
        "cm": "calc", 
        "bois": "check", 
        "beton": "calc", 
        "hydrodrain": "calc"
    }

    if output_format != "json":
        console.print(f"[bold cyan]--- Démarrage de l'analyse du projet ({generator.project_dir}) ---[/bold cyan]")
        if enable_cache:
            console.print("[cyan]Cache activé - Utilisation des résultats en cache si disponibles[/cyan]")
        console.print(f"[cyan]Parallélisation: {max_workers} workers[/cyan]")
    
    # Analyse parallèle avec cache
    results = generator.analyze_project_parallel(plugin_commands)

    if results:
        # Charger le rapport de comparaison si demandé
        comparison = None
        if compare_with and pathlib.Path(compare_with).exists():
            try:
                with open(compare_with, 'r', encoding='utf-8') as f:
                    previous_results = json.load(f)
                comparison = ReportAnalyzer.compare_reports(results, previous_results)
                console.print(f"[green]Comparaison chargée: {compare_with}[/green]")
            except Exception as e:
                console.print(f"[red]Erreur lors du chargement de la comparaison: {e}[/red]")
        
        # Générer la synthèse
        synthesis = ReportAnalyzer.generate_synthesis(results)
        
        if output_format == "json":
            json.dump(results, sys.stdout, indent=2, ensure_ascii=False)
            sys.stdout.write("\n")
        else:
            console.print(f"\n[bold cyan]--- Génération du rapport {output_format.upper()} ---[/bold cyan]")
            
            output_path = ""
            if output_format == "pdf":
                output_path = generator.generate_pdf_report(results, template_name, synthesis, comparison)
            elif output_format == "html":
                output_path = generator.generate_html_report(results, template_name, synthesis, comparison)
            elif output_format == "docx":
                output_path = generator.generate_docx_report(results, synthesis, comparison)
            elif output_format == "csv":
                output_path = generator.generate_csv_report(results)
            
            if output_path:
                console.print(f"[bold green]✓[/bold green] Rapport {output_format.upper()} généré : {output_path}")
                
                # Afficher les statistiques de synthèse
                console.print(f"\n[bold cyan]📊 Statistiques de synthèse:[/bold cyan]")
                console.print(f"  • Total éléments: {synthesis['total_elements']}")
                console.print(f"  • Taux de succès: {synthesis['success_rate']:.1f}%")
                console.print(f"  • Plugins utilisés: {len(synthesis['plugins_used'])}")
                console.print(f"  • Ratios critiques: {len(synthesis['critical_ratios'])}")
                
                if comparison:
                    console.print(f"\n[bold cyan]🔄 Comparaison:[/bold cyan]")
                    console.print(f"  • Éléments ajoutés: {comparison['summary_changes']['elements_added']}")
                    console.print(f"  • Éléments supprimés: {comparison['summary_changes']['elements_removed']}")
                    console.print(f"  • Éléments modifiés: {comparison['summary_changes']['elements_modified']}")
            else:
                console.print(f"[bold red]✗[/bold red] Échec de la génération du rapport {output_format.upper()}")
    else:
        if output_format != "json":
            console.print("\n[yellow]Aucune donnée d'analyse valide n'a été collectée. Le rapport n'a pas été généré.[/yellow]")