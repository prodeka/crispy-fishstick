# src/lcpi/reporting/utils/docx_generator.py

from pathlib import Path

try:
    import docx
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def export_to_docx(logs_data: list[dict], project_metadata: dict, output_path: Path):
    """
    Exporte les données de log dans un fichier DOCX.

    Note: C'est une implémentation de base. Pour une mise en forme avancée,
    l'utilisation d'un fichier .docx comme template est recommandée.

    Args:
        logs_data: La liste des données de log (dictionnaires Python).
        project_metadata: Les métadonnées du projet.
        output_path: Le chemin du fichier DOCX de sortie.
    """
    if not DOCX_AVAILABLE:
        print("Erreur : Le package 'python-docx' est requis pour l'export DOCX.")
        print("Veuillez l'installer avec : pip install python-docx")
        return

    try:
        document = docx.Document()

        # --- Style de base ---
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # --- En-tête ---
        document.add_heading('RAPPORT TECHNIQUE', level=0)
        document.add_heading(project_metadata.get("nom_projet", "Projet LCPI"), level=1)
        p = document.add_paragraph()
        p.add_run(f"Client: {project_metadata.get('client', 'N/A')}\n")
        p.add_run(f"Indice: {project_metadata.get('indice_revision', 'A')}")
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # --- Sections de calcul ---
        for i, log in enumerate(logs_data):
            document.add_heading(f"{i+1}. {log.get('titre_calcul', 'Calcul sans titre')}", level=2) 
            
            # Métadonnées du calcul
            meta_p = document.add_paragraph()
            meta_p.style = 'Intense Quote'
            meta_p.add_run(f"Date: {log.get('timestamp', 'N/A')}\n")
            meta_p.add_run(f"Commande: {log.get('commande_executee', 'N/A')}")

            document.add_heading('Résultats', level=3)

            # Contenu des résultats
            for key, value in log.get("donnees_resultat", {}).items():
                if isinstance(value, dict) and 'type_tableau' in value:
                    # C'est un tableau structuré
                    table_data = value
                    document.add_paragraph(table_data.get('titre', 'Tableau'), style='Caption')
                    
                    data_rows = table_data.get('donnees', [])
                    if not data_rows:
                        continue

                    headers = data_rows[0].keys()
                    table = document.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Shading Accent 1'
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header

                    for row_data in data_rows:
                        row_cells = table.add_row().cells
                        for i, key in enumerate(headers):
                            row_cells[i].text = str(row_data.get(key, ''))
                else:
                    # C'est une simple paire clé/valeur
                    document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}", style='List Bullet')

        document.save(output_path)
        print(f"✅ Rapport DOCX généré avec succès : {output_path}")

    except Exception as e:
        print(f"❌ Erreur lors de la génération du DOCX : {e}")
